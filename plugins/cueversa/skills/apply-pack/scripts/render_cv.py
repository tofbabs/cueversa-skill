#!/usr/bin/env python3
"""
Render a CV markdown file to .txt, .docx, and .pdf in one consistent template.

The .md is canonical. This script is deterministic formatting pushed into code,
so the rendered files never diverge from a hand-edited .docx — the same reason
scoring lives in score.py rather than in the model.

Dependencies, chosen so this runs in a code-execution sandbox with no system
binaries (no pandoc, no libreoffice, no weasyprint):
  - python-docx   (usually preinstalled)      -> .docx
  - fpdf2         (pip install fpdf2, pure Python) -> .pdf
  - .txt is a plain flatten, no dependency.

Usage:
  python3 render_cv.py CV.md
  python3 render_cv.py CV.md --out-dir /path/to/pack

Writes <stem>.txt, <stem>.docx, <stem>.pdf beside the input (or in --out-dir).

Supported CV markdown: `# Name`, `## Section`, `### Role`, `*italic meta*`,
`- bullet`, `**bold**` inline, and paragraphs. This is a CV renderer, not a
general markdown engine — it covers the shapes a CV uses and nothing more.
"""
from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

# --- a tiny, CV-shaped markdown model -------------------------------------


@dataclass
class Block:
    kind: str  # name | contact | section | role | meta | bullet | para
    text: str = ""
    items: list[str] = field(default_factory=list)


def parse(md: str) -> list[Block]:
    blocks: list[Block] = []
    lines = md.splitlines()
    i = 0
    seen_name = False
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith("### "):
            blocks.append(Block("role", line[4:].strip()))
        elif line.startswith("## "):
            blocks.append(Block("section", line[3:].strip()))
        elif line.startswith("# "):
            blocks.append(Block("name", line[2:].strip()))
            seen_name = True
        elif line.startswith("- ") or line.startswith("* "):
            items = []
            while i < len(lines) and lines[i].strip()[:2] in ("- ", "* "):
                items.append(lines[i].strip()[2:].strip())
                i += 1
            blocks.append(Block("bullet", items=items))
            continue
        elif line.startswith("*") and line.endswith("*") and len(line) > 2:
            blocks.append(Block("meta", line.strip("*").strip()))
        elif not seen_name is False and blocks and blocks[-1].kind == "name" and "·" in line:
            blocks.append(Block("contact", line))
        else:
            # First non-heading lines after the name, before any section, are contact.
            if blocks and blocks[-1].kind in ("name", "contact") and not any(
                b.kind == "section" for b in blocks
            ):
                blocks.append(Block("contact", line))
            else:
                blocks.append(Block("para", line))
        i += 1
    return blocks


# --- inline bold: **x** -> runs -------------------------------------------

_BOLD = re.compile(r"\*\*(.+?)\*\*")


def bold_spans(text: str):
    """Yield (chunk, is_bold) preserving **bold** segments."""
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            yield text[pos:m.start()], False
        yield m.group(1), True
        pos = m.end()
    if pos < len(text):
        yield text[pos:], False


def strip_md(text: str) -> str:
    return _BOLD.sub(r"\1", text)


# --- .txt ------------------------------------------------------------------


def render_txt(blocks: list[Block], out: Path) -> None:
    lines: list[str] = []
    for b in blocks:
        if b.kind == "name":
            lines += [strip_md(b.text).upper(), ""]
        elif b.kind == "contact":
            lines.append(strip_md(b.text))
        elif b.kind == "section":
            lines += ["", strip_md(b.text).upper(), "-" * len(b.text)]
        elif b.kind == "role":
            lines += ["", strip_md(b.text)]
        elif b.kind == "meta":
            lines.append(strip_md(b.text))
        elif b.kind == "bullet":
            lines += [f"  - {strip_md(it)}" for it in b.items]
        elif b.kind == "para":
            lines += ["", strip_md(b.text)]
    out.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


# --- .docx -----------------------------------------------------------------


def render_docx(blocks: list[Block], out: Path) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    accent = RGBColor(0x1D, 0x9E, 0x75)  # cueversa teal

    def add_runs(p, text):
        for chunk, is_bold in bold_spans(text):
            run = p.add_run(chunk)
            run.bold = is_bold

    for b in blocks:
        if b.kind == "name":
            p = doc.add_paragraph()
            r = p.add_run(strip_md(b.text))
            r.bold = True
            r.font.size = Pt(20)
            r.font.color.rgb = accent
        elif b.kind == "contact":
            p = doc.add_paragraph()
            r = p.add_run(strip_md(b.text))
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif b.kind == "section":
            p = doc.add_paragraph()
            p.space_before = Pt(10)
            r = p.add_run(strip_md(b.text).upper())
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = accent
        elif b.kind == "role":
            p = doc.add_paragraph()
            r = p.add_run(strip_md(b.text))
            r.bold = True
            r.font.size = Pt(10.5)
        elif b.kind == "meta":
            p = doc.add_paragraph()
            r = p.add_run(strip_md(b.text))
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif b.kind == "bullet":
            for it in b.items:
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, it)
        elif b.kind == "para":
            p = doc.add_paragraph()
            add_runs(p, b.text)
    doc.save(str(out))


# --- .pdf ------------------------------------------------------------------

# fpdf2 core fonts (Helvetica) cover only Latin-1, so an em-dash or a non-Latin
# name throws. We prefer a real Unicode TTF; if none is found we fall back to
# core fonts and normalise the common typographic characters to ASCII — which is
# also friendlier to ATS PDF parsers.

_SANITIZE = {
    "—": "-", "–": "-", "‒": "-", "−": "-",  # dashes
    "‘": "'", "’": "'", "“": '"', "”": '"',  # smart quotes
    "…": "...", " ": " ", "•": "-", "·": "-",
}


def _sanitize(text: str) -> str:
    for a, b in _SANITIZE.items():
        text = text.replace(a, b)
    # drop anything still outside Latin-1 rather than crash
    return text.encode("latin-1", "replace").decode("latin-1")


def _find_unicode_font() -> tuple[str, str] | None:
    """Return (regular, bold) TTF paths if a Unicode font is available.

    Order: a font bundled in the skill's assets/ (deterministic, preferred for a
    distributed skill); then matplotlib's DejaVuSans (present in the code-exec
    sandbox); then common OS locations. Returns None to trigger the ASCII fallback.
    """
    import glob
    import os

    here = Path(__file__).resolve().parent
    bundled = here.parent / "assets" / "DejaVuSans.ttf"
    if bundled.exists():
        b = bundled.with_name("DejaVuSans-Bold.ttf")
        return str(bundled), str(b if b.exists() else bundled)

    try:
        import matplotlib

        d = os.path.join(os.path.dirname(matplotlib.__file__), "mpl-data", "fonts", "ttf")
        reg = os.path.join(d, "DejaVuSans.ttf")
        if os.path.exists(reg):
            bold = os.path.join(d, "DejaVuSans-Bold.ttf")
            return reg, bold if os.path.exists(bold) else reg
    except Exception:
        pass

    # Only DejaVu is trusted. Arial Unicode (macOS) triggers an fpdf2 width bug
    # when a bold heading precedes body text, so it is deliberately excluded —
    # a missing DejaVu takes the ASCII sanitize path, which is well-behaved.
    dj = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    if os.path.exists(dj):
        bold = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
        return dj, bold if os.path.exists(bold) else dj
    for reg in glob.glob("/usr/share/fonts/**/DejaVuSans.ttf", recursive=True):
        return reg, reg
    return None


def render_pdf(blocks: list[Block], out: Path) -> None:
    from fpdf import FPDF

    TEAL = (0x1D, 0x9E, 0x75)
    GREY = (0x55, 0x55, 0x55)
    BLACK = (0x1A, 0x1A, 0x1A)

    font_paths = _find_unicode_font()
    FAM = "Body"

    pdf = FPDF(format="A4", unit="mm")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(18, 16, 18)

    if font_paths:
        reg, bold = font_paths
        pdf.add_font(FAM, "", reg)
        pdf.add_font(FAM, "B", bold)
        pdf.add_font(FAM, "I", reg)  # DejaVu italic optional; regular is fine
        prep = lambda s: s  # noqa: E731  Unicode font: pass text through
    else:
        FAM = "Helvetica"
        prep = _sanitize

    def cell(text, size, style="", color=BLACK, h=5.0):
        pdf.set_text_color(*color)
        pdf.set_font(FAM, style, size)
        # new_x=LMARGIN: fpdf2 otherwise leaves the cursor at the right edge, so
        # the next full-width multi_cell would see ~0 available width and throw.
        pdf.multi_cell(0, h, prep(text), new_x="LMARGIN", new_y="NEXT")

    for b in blocks:
        if b.kind == "name":
            cell(strip_md(b.text), 20, "B", TEAL, 9)
        elif b.kind == "contact":
            cell(strip_md(b.text), 8.5, "", GREY, 4.5)
        elif b.kind == "section":
            pdf.ln(3)
            cell(strip_md(b.text).upper(), 11, "B", TEAL, 6)
            pdf.set_draw_color(*TEAL)
            y = pdf.get_y()
            pdf.line(18, y, 192, y)
            pdf.ln(1.5)
        elif b.kind == "role":
            cell(strip_md(b.text), 10.5, "B", BLACK, 5.5)
        elif b.kind == "meta":
            cell(strip_md(b.text), 8.5, "I", GREY, 4.5)
        elif b.kind == "bullet":
            for it in b.items:
                cell("- " + strip_md(it), 10, "", BLACK, 5)
        elif b.kind == "para":
            cell(b.text, 10, "", BLACK, 5)
    pdf.output(str(out))


# --- cli -------------------------------------------------------------------


def main(argv: list[str]) -> int:
    ap = argparse.ArgumentParser(description="Render a CV markdown to txt/docx/pdf.")
    ap.add_argument("md", help="path to the canonical CV markdown")
    ap.add_argument("--out-dir", help="output directory (defaults beside the input)")
    args = ap.parse_args(argv)

    src = Path(args.md)
    if not src.exists():
        print(f"not found: {src}", file=sys.stderr)
        return 1
    out_dir = Path(args.out_dir) if args.out_dir else src.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = src.stem

    blocks = parse(src.read_text(encoding="utf-8"))
    render_txt(blocks, out_dir / f"{stem}.txt")
    render_docx(blocks, out_dir / f"{stem}.docx")
    render_pdf(blocks, out_dir / f"{stem}.pdf")
    for ext in ("txt", "docx", "pdf"):
        print(f"wrote {out_dir / f'{stem}.{ext}'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
